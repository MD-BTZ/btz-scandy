from flask import Blueprint, render_template, request, jsonify, session, redirect, url_for, flash, abort, send_file, render_template_string, current_app
from app.models.mongodb_models import MongoDBTicket
from app.models.mongodb_database import mongodb
from app.utils.decorators import login_required, admin_required, not_teilnehmer_required
from app.utils.database_helpers import get_ticket_categories_from_settings, get_categories_from_settings, get_next_ticket_number
from app.models.user import User
import logging
from datetime import datetime
from flask_login import current_user
from docxtpl import DocxTemplate
import os
from bson import ObjectId

bp = Blueprint('tickets', __name__, url_prefix='/tickets')

@bp.route('/')
@login_required
@not_teilnehmer_required
def index():
    """Zeigt die Ticket-Übersicht für den Benutzer."""
    # Hole die eigenen Tickets für die Anzeige (nur nicht geschlossene und nicht gelöschte)
    my_tickets = mongodb.find('tickets', {
        'created_by': current_user.username,
        'status': {'$ne': 'geschlossen'},
        'deleted': {'$ne': True}
    })
    my_tickets = list(my_tickets)
    
    # Füge Nachrichtenanzahl hinzu
    for ticket in my_tickets:
        message_count = mongodb.count_documents('ticket_messages', {'ticket_id': ticket['_id']})
        ticket['message_count'] = message_count
    
    # Sortiere nach Erstellungsdatum (neueste zuerst)
    my_tickets.sort(key=lambda x: x.get('created_at', datetime.min), reverse=True)
            
    # Füge id-Feld zu allen Tickets hinzu (für Template-Kompatibilität)
    for ticket in my_tickets:
        ticket['id'] = str(ticket['_id'])
    
    return render_template('tickets/index.html', tickets=my_tickets)

@bp.route('/create', methods=['GET', 'POST'])
@login_required
def create():
    """Erstellt ein neues Ticket."""
    if request.method == 'POST':
        try:
            # Prüfe ob die Anfrage JSON enthält
            if request.is_json:
                data = request.get_json()
                title = data.get('title')
                description = data.get('description')
                priority = data.get('priority', 'normal')
                category = data.get('category')
                new_category = data.get('new_category')
            else:
                title = request.form.get('title')
                description = request.form.get('description')
                priority = request.form.get('priority', 'normal')
                category = request.form.get('category')
                new_category = request.form.get('new_category')

            # Wenn eine neue Kategorie eingegeben wurde, prüfe und füge sie ggf. hinzu
            if category:
                # Prüfe ob die Kategorie bereits in den Settings existiert
                ticket_categories = get_ticket_categories_from_settings()
                if category not in ticket_categories:
                    # Füge die neue Kategorie zu den Settings hinzu
                    mongodb.update_one_array(
                        'settings',
                        {'key': 'ticket_categories'},
                        {'$push': {'value': category}},
                        upsert=True
                    )

            due_date = data.get('due_date') if request.is_json else request.form.get('due_date')
            estimated_time = data.get('estimated_time') if request.is_json else request.form.get('estimated_time')

            # Formatiere das Fälligkeitsdatum
            if due_date:
                try:
                    due_date = datetime.strptime(due_date, '%Y-%m-%dT%H:%M')
                except ValueError:
                    due_date = None

            if not title:
                if request.is_json:
                    return jsonify({'success': False, 'message': 'Titel ist erforderlich'}), 400
                flash('Titel ist erforderlich.', 'error')
                return redirect(url_for('tickets.create'))

            # Erstelle das Ticket
            ticket_data = {
                'title': title,
                'description': description,
                'priority': priority,
                'created_by': current_user.username,
                'category': category,
                'due_date': due_date,
                'estimated_time': estimated_time,
                'status': 'offen',
                'created_at': datetime.now(),
                'updated_at': datetime.now(),
                'ticket_number': get_next_ticket_number()  # Neue Auftragsnummer
            }
            
            result = mongodb.insert_one('tickets', ticket_data)
            ticket_id = str(result)

            if request.is_json:
                return jsonify({
                    'success': True,
                    'message': 'Ticket wurde erfolgreich erstellt',
                    'ticket_id': ticket_id
                })
            
            flash('Ticket wurde erfolgreich erstellt.', 'success')
            return redirect(url_for('tickets.create'))
        except Exception as e:
            if request.is_json:
                return jsonify({
                    'success': False,
                    'message': f'Fehler beim Erstellen des Tickets: {str(e)}'
                }), 500
            flash(f'Fehler beim Erstellen des Tickets: {str(e)}', 'error')
            return redirect(url_for('tickets.create'))
    
    # Hole die eigenen Tickets für die Anzeige (nur nicht geschlossene und nicht gelöschte)
    my_tickets = mongodb.find('tickets', {
        'created_by': current_user.username,
        'status': {'$ne': 'geschlossen'},
        'deleted': {'$ne': True}
    })
    my_tickets = list(my_tickets)
    
    # Füge Nachrichtenanzahl hinzu
    for ticket in my_tickets:
        message_count = mongodb.count_documents('ticket_messages', {'ticket_id': ticket['_id']})
        ticket['message_count'] = message_count

    # Hole die zugewiesenen Tickets (nur nicht geschlossene und nicht gelöschte)
    assigned_tickets = mongodb.find('tickets', {
        'assigned_to': current_user.username,
        'status': {'$ne': 'geschlossen'},
        'deleted': {'$ne': True}
    })
    assigned_tickets = list(assigned_tickets)
    
    # Füge Nachrichtenanzahl hinzu
    for ticket in assigned_tickets:
        message_count = mongodb.count_documents('ticket_messages', {'ticket_id': ticket['_id']})
        ticket['message_count'] = message_count

    # Hole offene Tickets (nur nicht geschlossene und nicht gelöschte)
    open_tickets = mongodb.find('tickets', {
        '$or': [
            {'assigned_to': None},
            {'assigned_to': ''}
        ],
        'status': 'offen',
        'deleted': {'$ne': True}
    })
    open_tickets = list(open_tickets)
    
    # Füge Nachrichtenanzahl hinzu
    for ticket in open_tickets:
        message_count = mongodb.count_documents('ticket_messages', {'ticket_id': ticket['_id']})
        ticket['message_count'] = message_count
    
    # Hole alle Kategorien aus der settings Collection
    categories = get_ticket_categories_from_settings()
    
    # Sortiere alle Listen nach Erstellungsdatum (neueste zuerst)
    my_tickets.sort(key=lambda x: x.get('created_at', datetime.min), reverse=True)
    assigned_tickets.sort(key=lambda x: x.get('created_at', datetime.min), reverse=True)
    open_tickets.sort(key=lambda x: x.get('created_at', datetime.min), reverse=True)
    
    # Füge id-Feld zu allen Tickets hinzu (für Template-Kompatibilität)
    for ticket in my_tickets:
        ticket['id'] = str(ticket['_id'])
    for ticket in assigned_tickets:
        ticket['id'] = str(ticket['_id'])
    for ticket in open_tickets:
        ticket['id'] = str(ticket['_id'])
            
    return render_template('tickets/create.html', 
                         my_tickets=my_tickets,
                         assigned_tickets=assigned_tickets,
                         open_tickets=open_tickets,
                         categories=categories,
                         status_colors={
                             'offen': 'info',
                             'in_bearbeitung': 'warning',
                             'wartet_auf_antwort': 'warning',
                             'gelöst': 'success',
                             'geschlossen': 'ghost'
                         },
                         priority_colors={
                             'niedrig': 'secondary',
                             'normal': 'primary',
                             'hoch': 'error',
                             'dringend': 'error'
                         })

@bp.route('/view/<ticket_id>')
@login_required
def view(ticket_id):
    """Zeigt die Details eines Tickets für den Benutzer."""
    logging.info(f"Lade Ticket {ticket_id} für Benutzer {current_user.username}")
    
    # Konvertiere ticket_id zu ObjectId
    try:
        object_id = ObjectId(ticket_id)
    except:
        flash('Ungültige Ticket-ID.', 'error')
        return redirect(url_for('tickets.create'))
    
    ticket = mongodb.find_one('tickets', {'_id': object_id})
    
    if not ticket:
        logging.error(f"Ticket {ticket_id} nicht gefunden")
        flash('Ticket nicht gefunden.', 'error')
        return redirect(url_for('tickets.create'))
        
    # Prüfe ob der Benutzer berechtigt ist, das Ticket zu sehen
    if (
        ticket.get('created_by') != current_user.username
        and ticket.get('assigned_to') not in [None, '', current_user.username]
        and current_user.role not in ['admin', 'mitarbeiter']
    ):
        logging.error(f"Benutzer {current_user.username} hat keine Berechtigung für Ticket {ticket_id}")
        flash('Sie haben keine Berechtigung, dieses Ticket zu sehen.', 'error')
        return redirect(url_for('tickets.create'))
    
    # Hole die Nachrichten für das Ticket
    logging.info(f"Hole Nachrichten für Ticket {ticket_id}")
    
    try:
        messages = mongodb.find('ticket_messages', {'ticket_id': object_id})
        messages = list(messages)
        
        # Sortiere Nachrichten nach Datum (älteste zuerst)
        messages.sort(key=lambda x: x.get('created_at', datetime.min))
        
        # Formatiere Datum für jede Nachricht
        for msg in messages:
            if isinstance(msg.get('created_at'), datetime):
                msg['formatted_date'] = msg['created_at'].strftime('%d.%m.%Y %H:%M')
            else:
                msg['formatted_date'] = str(msg.get('created_at', ''))
        
        logging.info(f"Nachrichtenabfrage ergab {len(messages)} Nachrichten")
        
        # Füge id-Feld hinzu (für Template-Kompatibilität)
        ticket['id'] = str(ticket['_id'])
        
        # Hole die Auftragsdetails
        auftrag_details = mongodb.find_one('auftrag_details', {'ticket_id': object_id})
        
        # Hole Materialliste
        material_list = list(mongodb.find('auftrag_material', {'ticket_id': object_id}))
        
        # Hole Arbeitsliste
        arbeit_list = list(mongodb.find('auftrag_arbeit', {'ticket_id': object_id}))
        
        # Hole alle Kategorien aus der settings Collection
        categories = get_ticket_categories_from_settings()
        
        # Hole alle Benutzer für die Zuweisung (falls benötigt)
        users = mongodb.find('users', {'is_active': True})
        users = [dict(user) for user in users]
        
        return render_template('tickets/view.html', 
                             ticket=ticket, 
                             messages=messages,
                             auftrag_details=auftrag_details,
                             categories=categories,
                             workers=users,
                             status_colors={
                                 'offen': 'info',
                                 'in_bearbeitung': 'warning',
                                 'wartet_auf_antwort': 'warning',
                                 'gelöst': 'success',
                                 'geschlossen': 'ghost'
                             })
                             
    except Exception as e:
        logging.error(f"Fehler beim Laden der Nachrichten: {str(e)}")
        flash('Fehler beim Laden der Nachrichten.', 'error')
        return redirect(url_for('tickets.create'))

@bp.route('/<ticket_id>/message', methods=['POST'])
@login_required
def add_message(ticket_id):
    """Fügt eine neue Nachricht zu einem Ticket hinzu"""
    try:
        # Prüfe ob das Ticket existiert
        ticket = mongodb.find_one('tickets', {'_id': ObjectId(ticket_id)})
        if not ticket:
            logging.error(f"Ticket {ticket_id} nicht gefunden")
            return jsonify({'success': False, 'message': 'Ticket nicht gefunden'}), 404
        
        # Hole die Nachricht aus dem Request
        if not request.is_json:
            logging.error("Ungültiges Anfrageformat (kein JSON)")
            return jsonify({'success': False, 'message': 'Ungültiges Anfrageformat'}), 400
            
        data = request.get_json()
        message = data['message'].strip()
        if not message:
            logging.error("Leere Nachricht")
            return jsonify({'success': False, 'message': 'Nachricht darf nicht leer sein'}), 400

        logging.info(f"Versuche Nachricht zu speichern: Ticket {ticket_id}, Benutzer {current_user.username}, Nachricht: {message}")

        # Verwende mongodb.insert_one für die Nachrichtenspeicherung
        message_data = {
            'ticket_id': ObjectId(ticket_id),
            'message': message,
            'sender': current_user.username,
            'created_at': datetime.now()
        }
        
        result = mongodb.insert_one('ticket_messages', message_data)
        
        logging.info(f"Nachricht erfolgreich gespeichert")

        # Hole die aktuelle Zeit für die Antwort
        created_at = datetime.now()
        formatted_date = created_at.strftime('%d.%m.%Y, %H:%M')

        return jsonify({
            'success': True,
            'message': {
                'text': message,
                'sender': current_user.username,
                'created_at': formatted_date
            }
        })

    except Exception as e:
        logging.error(f"Fehler beim Hinzufügen der Nachricht: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': 'Ein Fehler ist aufgetreten'}), 500

@bp.route('/<id>')
@login_required
def detail(id):
    """Zeigt die Details eines Tickets."""
    ticket = mongodb.find_one('tickets', {'_id': ObjectId(id)})
    
    if not ticket:
        return render_template('404.html'), 404
    
    # Prüfe Berechtigungen: Normale User können nur ihre eigenen Tickets sehen
    if current_user.role not in ['admin', 'mitarbeiter'] and ticket.get('created_by') != current_user.username:
        flash('Sie haben keine Berechtigung, dieses Ticket zu sehen.', 'error')
        return redirect(url_for('tickets.index'))
    
    # Füge id-Feld hinzu (für Template-Kompatibilität)
    ticket['id'] = str(ticket['_id'])
    
    # Konvertiere alle Datumsfelder zu datetime-Objekten
    date_fields = ['created_at', 'updated_at', 'resolved_at', 'due_date']
    for field in date_fields:
        if ticket.get(field):
            try:
                ticket[field] = datetime.strptime(ticket[field], '%Y-%m-%d %H:%M:%S')
            except (ValueError, TypeError):
                ticket[field] = None
        
    # Hole die Notizen für das Ticket
    notes = mongodb.find('ticket_notes', {'ticket_id': ObjectId(id)})

    # Hole die Nachrichten für das Ticket
    messages = mongodb.find('ticket_messages', {'ticket_id': ObjectId(id)})

    # Hole die Auftragsdetails
    auftrag_details = mongodb.find_one('auftrag_details', {'ticket_id': ObjectId(id)})
    
    # Hole Materialliste
    material_list = list(mongodb.find('auftrag_material', {'ticket_id': ObjectId(id)}))
    
    # Hole Arbeitsliste
    arbeit_list = list(mongodb.find('auftrag_arbeit', {'ticket_id': ObjectId(id)}))

    # Hole alle Benutzer aus der Hauptdatenbank und wandle sie in Dicts um
    users = mongodb.find('users', {'is_active': True})
    users = [dict(user) for user in users]

    # Hole alle zugewiesenen Nutzer (Mehrfachzuweisung)
    assigned_users = mongodb.find('ticket_assignments', {'ticket_id': ObjectId(id)})

    # Hole alle Kategorien aus der settings Collection
    categories = get_ticket_categories_from_settings()

    # Bestimme Berechtigungen
    can_edit = current_user.role in ['admin', 'mitarbeiter'] or ticket.get('created_by') == current_user.username
    can_assign = current_user.role in ['admin', 'mitarbeiter']
    can_change_status = current_user.role in ['admin', 'mitarbeiter']
    can_delete = current_user.role in ['admin', 'mitarbeiter']

    return render_template('tickets/detail.html', 
                         ticket=ticket, 
                         notes=notes,
                         messages=messages,
                         users=users,
                         assigned_users=assigned_users,
                         auftrag_details=auftrag_details,
                         material_list=material_list,
                         categories=categories,
                         can_edit=can_edit,
                         can_assign=can_assign,
                         can_change_status=can_change_status,
                         can_delete=can_delete,
                         now=datetime.now())

@bp.route('/<id>/delete', methods=['POST'])
@admin_required
def delete(id):
    """Löscht ein Ticket."""
    try:
        # Prüfe ob Ticket existiert
        ticket = mongodb.find_one('tickets', {'_id': ObjectId(id)})
        
        if not ticket:
            return jsonify({
                'success': False,
                'message': 'Ticket nicht gefunden'
            }), 404
            
        # Lösche das Ticket
        if not mongodb.delete_one('tickets', {'_id': ObjectId(id)}):
            return jsonify({
                'success': False,
                'message': 'Fehler beim Löschen des Tickets'
            }), 500
        
        return jsonify({
            'success': True,
            'message': 'Ticket wurde gelöscht'
        })
        
    except Exception as e:
        logging.error(f"Fehler beim Löschen des Tickets #{id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': 'Fehler beim Löschen des Tickets'
        }), 500

@bp.route('/<id>/auftrag-details-modal')
@login_required
def auftrag_details_modal(id):
    ticket = mongodb.find_one('tickets', {'_id': ObjectId(id)})
    
    if not ticket:
        return render_template('404.html'), 404
    
    # Füge id-Feld hinzu (für Template-Kompatibilität)
    ticket['id'] = str(ticket['_id'])
    
    # Konvertiere alle Datumsfelder zu datetime-Objekten
    date_fields = ['created_at', 'updated_at', 'resolved_at', 'due_date']
    for field in date_fields:
        if ticket.get(field):
            try:
                ticket[field] = datetime.strptime(ticket[field], '%Y-%m-%d %H:%M:%S')
            except (ValueError, TypeError):
                ticket[field] = None
        
    # Hole die Notizen für das Ticket
    notes = mongodb.find('ticket_notes', {'ticket_id': ObjectId(id)})

    # Hole die Nachrichten für das Ticket
    messages = mongodb.find('ticket_messages', {'ticket_id': ObjectId(id)})

    # Hole die Auftragsdetails
    auftrag_details = mongodb.find_one('auftrag_details', {'ticket_id': ObjectId(id)})
    
    # Hole Materialliste
    material_list = list(mongodb.find('auftrag_material', {'ticket_id': ObjectId(id)}))
    
    # Hole Arbeitsliste
    arbeit_list = list(mongodb.find('auftrag_arbeit', {'ticket_id': ObjectId(id)}))

    return render_template('tickets/auftrag_details_modal.html', 
                         ticket=ticket, 
                         notes=notes,
                         messages=messages,
                         auftrag_details=auftrag_details,
                         material_list=material_list,
                         arbeit_list=arbeit_list)

@bp.route('/<id>/update-status', methods=['POST'])
@login_required
def update_status(id):
    """Aktualisiert den Status eines Tickets und weist es ggf. automatisch zu"""
    try:
        if not request.is_json:
            return jsonify({'success': False, 'message': 'Ungültiges Anfrageformat'}), 400

        data = request.get_json()
        new_status = data.get('status')
        if not new_status:
            return jsonify({'success': False, 'message': 'Status ist erforderlich'}), 400

        ticket = mongodb.find_one('tickets', {'_id': ObjectId(id)})
        if not ticket:
            return jsonify({'success': False, 'message': 'Ticket nicht gefunden'}), 404

        update_fields = {'status': new_status, 'updated_at': datetime.now()}

        # Automatische Zuweisung: Wenn Status von 'offen' auf etwas anderes wechselt und noch niemand zugewiesen ist
        if new_status != 'offen' and (not ticket.get('assigned_to')):
            update_fields['assigned_to'] = current_user.username
        # Wenn Status auf 'offen' gesetzt wird, Zuweisung entfernen
        elif new_status == 'offen':
            update_fields['assigned_to'] = None

        if not mongodb.update_one('tickets', {'_id': ObjectId(id)}, {'$set': update_fields}):
            return jsonify({'success': False, 'message': 'Fehler beim Aktualisieren des Status'})

        # Spezielle Nachricht bei automatischer Zuweisung
        if new_status != 'offen' and (not ticket.get('assigned_to')):
            return jsonify({
                'success': True, 
                'message': f'Status erfolgreich auf "{new_status}" geändert. Ticket wurde automatisch Ihnen zugewiesen.'
            })
        elif new_status == 'offen':
            return jsonify({
                'success': True, 
                'message': f'Status erfolgreich auf "{new_status}" geändert. Zuweisung wurde entfernt.'
            })
        else:
            return jsonify({'success': True, 'message': 'Status erfolgreich aktualisiert'})

    except Exception as e:
        logging.error(f"Fehler beim Aktualisieren des Status: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@bp.route('/<id>/update-assignment', methods=['POST'])
@login_required
def update_assignment(id):
    """Aktualisiert die Zuweisung eines Tickets"""
    try:
        if not request.is_json:
            return jsonify({'success': False, 'message': 'Ungültiges Anfrageformat'}), 400

        data = request.get_json()
        assigned_to = data.get('assigned_to')
        
        # Wenn assigned_to leer ist, setze es auf None
        if not assigned_to or assigned_to == "":
            assigned_to = None

        # Aktualisiere die Zuweisung direkt im Ticket
        if not mongodb.update_one('tickets', {'_id': ObjectId(id)}, {'$set': {'assigned_to': assigned_to, 'updated_at': datetime.now()}}):
            return jsonify({'success': False, 'message': 'Fehler beim Aktualisieren der Zuweisung'})

        return jsonify({'success': True, 'message': 'Zuweisung erfolgreich aktualisiert'})

    except Exception as e:
        logging.error(f"Fehler beim Aktualisieren der Zuweisung: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

@bp.route('/<id>/update-details', methods=['POST'])
@login_required
def update_details(id):
    """Aktualisiert die Auftragsdetails eines Tickets"""
    try:
        # Prüfe ob Ticket existiert
        ticket = mongodb.find_one('tickets', {'_id': ObjectId(id)})
        if not ticket:
            if request.is_json:
                return jsonify({'success': False, 'message': 'Ticket nicht gefunden'}), 404
            else:
                flash('Ticket nicht gefunden', 'error')
                return redirect(url_for('tickets.index'))
        
        # Verarbeite die Daten je nach Request-Typ
        if request.is_json:
            data = request.get_json()
            if not data:
                return jsonify({'success': False, 'message': 'Ungültiges Anfrageformat'}), 400
        else:
            # Formular-Daten verarbeiten
            form_data = request.form.to_dict()
            data = {}
            
            # Basis-Formulardaten
            for key, value in form_data.items():
                data[key] = value
            
            # Verarbeite den Auftraggeber-Typ (Radio-Button)
            auftraggeber_typ = data.get('auftraggeber_typ', '')
            data['auftraggeber_intern'] = auftraggeber_typ == 'intern'
            data['auftraggeber_extern'] = auftraggeber_typ == 'extern'
            
            # Materialliste aus Formular sammeln
            material_list = []
            material_names = request.form.getlist('material')
            material_mengen = request.form.getlist('menge')
            material_preise = request.form.getlist('einzelpreis')
            
            for i in range(len(material_names)):
                if material_names[i] or material_mengen[i] or material_preise[i]:
                    material_list.append({
                        'material': material_names[i],
                        'menge': float(material_mengen[i]) if material_mengen[i] else 0,
                        'einzelpreis': float(material_preise[i]) if material_preise[i] else 0
                    })
            
            # Arbeitsliste aus Formular sammeln
            arbeit_list = []
            arbeit_names = request.form.getlist('arbeit')
            arbeit_stunden = request.form.getlist('arbeitsstunden')
            arbeit_kategorien = request.form.getlist('leistungskategorie')
            
            for i in range(len(arbeit_names)):
                if arbeit_names[i] or arbeit_stunden[i] or arbeit_kategorien[i]:
                    arbeit_list.append({
                        'arbeit': arbeit_names[i],
                        'arbeitsstunden': float(arbeit_stunden[i]) if arbeit_stunden[i] else 0,
                        'leistungskategorie': arbeit_kategorien[i]
                    })
            
            data['material_list'] = material_list
            data['arbeit_list'] = arbeit_list

        logging.info(f"Empfangene Daten für Ticket {id}: {data}")
        
        # Verarbeite ausgeführte Arbeiten
        arbeit_list = data.get('arbeit_list', [])
        ausgefuehrte_arbeiten = '\n'.join([
            f"{arbeit.get('arbeit', '')}|{arbeit.get('arbeitsstunden', '')}|{arbeit.get('leistungskategorie', '')}"
            for arbeit in arbeit_list
        ])
        logging.info(f"Verarbeitete ausgeführte Arbeiten: {ausgefuehrte_arbeiten}")
        
        # Bereite die Auftragsdetails vor
        auftrag_details_daten = {
            'ticket_id': ObjectId(id),
            'bereich': data.get('bereich', ''),
            'auftraggeber_intern': data.get('auftraggeber_intern', False),
            'auftraggeber_extern': data.get('auftraggeber_extern', False),
            'auftraggeber_name': data.get('auftraggeber_name', ''),
            'kontakt': data.get('kontakt', ''),
            'auftragsbeschreibung': data.get('auftragsbeschreibung', ''),
            'ausgefuehrte_arbeiten': ausgefuehrte_arbeiten,
            'fertigstellungstermin': data.get('fertigstellungstermin'),
            'gesamtsumme': data.get('gesamtsumme', 0)
        }
        
        # Speichere oder aktualisiere die Auftragsdetails
        existing_details = mongodb.find_one('auftrag_details', {'ticket_id': ObjectId(id)})
        if existing_details:
            mongodb.update_one('auftrag_details', {'ticket_id': ObjectId(id)}, {'$set': auftrag_details_daten})
        else:
            mongodb.insert_one('auftrag_details', auftrag_details_daten)
        
        # Verarbeite die Materialliste
        material_list = data.get('material_list', [])
        mongodb.delete_many('auftrag_material', {'ticket_id': ObjectId(id)})
        
        if material_list:
            material_daten = [{
                'ticket_id': ObjectId(id),
                'material': m.get('material', ''),
                'menge': m.get('menge', 0),
                'einzelpreis': m.get('einzelpreis', 0)
            } for m in material_list]
            mongodb.insert_many('auftrag_material', material_daten)
        
        # Setze das 'updated_at' Feld am Ticket selbst
        mongodb.update_one('tickets', {'_id': ObjectId(id)}, {'$set': {'updated_at': datetime.now()}})

        # Rückgabe je nach Request-Typ
        if request.is_json:
            return jsonify({'success': True, 'message': 'Auftragsdetails erfolgreich gespeichert'})
        else:
            flash('Auftragsdetails erfolgreich gespeichert', 'success')
            return redirect(url_for('tickets.view', ticket_id=id))

    except Exception as e:
        logging.error(f"Fehler beim Aktualisieren der Auftragsdetails: {e}", exc_info=True)
        if request.is_json:
            return jsonify({'success': False, 'message': str(e)}), 500
        else:
            flash(f'Fehler beim Speichern: {str(e)}', 'error')
            return redirect(url_for('tickets.auftrag_details_page', id=id))

@bp.route('/<id>/export')
@login_required
@admin_required
def export_ticket(id):
    """Exportiert ein Ticket als Word-Dokument"""
    try:
        # Hole Ticket-Daten
        ticket = mongodb.find_one('tickets', {'_id': ObjectId(id)})
        if not ticket:
            flash('Ticket nicht gefunden', 'error')
            return redirect(url_for('tickets.index'))
        
        # Hole Auftragsdetails
        auftrag_details = mongodb.find_one('auftrag_details', {'ticket_id': ObjectId(id)})
        
        # Hole Materialliste
        material_list = list(mongodb.find('auftrag_material', {'ticket_id': ObjectId(id)}))
        
        # Hole Arbeitszeiten
        arbeit_list = list(mongodb.find('auftrag_arbeit', {'ticket_id': ObjectId(id)}))
        
        # Erstelle Word-Dokument
        from docx import Document
        from docx.shared import Inches
        import tempfile
        import os
        
        doc = Document()
        
        # Titel
        title = doc.add_heading(f'Auftrag: {ticket.get("title", "Unbekannt")}', 0)
        title.alignment = 1  # Zentriert
        
        # Ticket-Informationen
        doc.add_heading('Ticket-Informationen', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header-Zeile
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Feld'
        hdr_cells[1].text = 'Wert'
        
        # Ticket-Daten
        ticket_data = [
            ('Ticket-ID', str(ticket.get('_id', ''))),
            ('Titel', ticket.get('title', '')),
            ('Beschreibung', ticket.get('description', '')),
            ('Status', ticket.get('status', '')),
            ('Priorität', ticket.get('priority', '')),
            ('Erstellt von', ticket.get('created_by', '')),
            ('Erstellt am', ticket.get('created_at', '').strftime('%d.%m.%Y %H:%M') if ticket.get('created_at') else ''),
            ('Zugewiesen an', ticket.get('assigned_to', '')),
            ('Kategorie', ticket.get('category', ''))
        ]
        
        for field, value in ticket_data:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = str(value)
        
        # Auftragsdetails
        if auftrag_details:
            doc.add_heading('Auftragsdetails', level=1)
            auftrag_table = doc.add_table(rows=1, cols=2)
            auftrag_table.style = 'Table Grid'
            
            # Header
            auftrag_hdr = auftrag_table.rows[0].cells
            auftrag_hdr[0].text = 'Feld'
            auftrag_hdr[1].text = 'Wert'
            
            # Auftragsdaten
            auftrag_data = [
                ('Kunde', auftrag_details.get('customer_name', '')),
                ('Kontakt', auftrag_details.get('customer_contact', '')),
                ('Bereich', auftrag_details.get('customer_department', '')),
                ('Auftragsnummer', auftrag_details.get('order_number', '')),
                ('Startdatum', auftrag_details.get('start_date', '').strftime('%d.%m.%Y') if auftrag_details.get('start_date') else ''),
                ('Enddatum', auftrag_details.get('end_date', '').strftime('%d.%m.%Y') if auftrag_details.get('end_date') else '')
            ]
            
            for field, value in auftrag_data:
                row_cells = auftrag_table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
        
        # Materialliste
        if material_list:
            doc.add_heading('Materialliste', level=1)
            material_table = doc.add_table(rows=1, cols=4)
            material_table.style = 'Table Grid'
            
            # Header
            material_hdr = material_table.rows[0].cells
            material_hdr[0].text = 'Material'
            material_hdr[1].text = 'Menge'
            material_hdr[2].text = 'Einzelpreis'
            material_hdr[3].text = 'Gesamtpreis'
            
            # Materialdaten
            for material in material_list:
                row_cells = material_table.add_row().cells
                row_cells[0].text = material.get('name', '')
                row_cells[1].text = str(material.get('quantity', 0))
                row_cells[2].text = f"{material.get('unit_price', 0):.2f} €"
                row_cells[3].text = f"{material.get('total_price', 0):.2f} €"
        
        # Arbeitszeiten
        if arbeit_list:
            doc.add_heading('Arbeitszeiten', level=1)
            arbeit_table = doc.add_table(rows=1, cols=4)
            arbeit_table.style = 'Table Grid'
            
            # Header
            arbeit_hdr = arbeit_table.rows[0].cells
            arbeit_hdr[0].text = 'Beschreibung'
            arbeit_hdr[1].text = 'Stunden'
            arbeit_hdr[2].text = 'Stundensatz'
            arbeit_hdr[3].text = 'Gesamtpreis'
            
            # Arbeitsdaten
            for arbeit in arbeit_list:
                row_cells = arbeit_table.add_row().cells
                row_cells[0].text = arbeit.get('description', '')
                row_cells[1].text = str(arbeit.get('hours', 0))
                row_cells[2].text = f"{arbeit.get('hourly_rate', 0):.2f} €"
                row_cells[3].text = f"{arbeit.get('total_price', 0):.2f} €"
        
        # Speichere temporäre Datei
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_file_path = tmp_file.name
        
        # Sende Datei als Download
        return send_file(
            tmp_file_path,
            as_attachment=True,
            download_name=f'ticket_{id}_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logging.error(f"Fehler beim Exportieren des Tickets: {str(e)}", exc_info=True)
        flash('Fehler beim Exportieren des Tickets', 'error')
        return redirect(url_for('tickets.index'))

@bp.route('/<id>/note', methods=['POST'])
@login_required
@admin_required
def add_note(id):
    """Fügt eine neue Notiz zu einem Ticket hinzu"""
    try:
        if not request.is_json:
            return jsonify({'success': False, 'message': 'Ungültiges Anfrageformat'}), 400

        data = request.get_json()
        note_text = data.get('note', '').strip()
        
        if not note_text:
            return jsonify({'success': False, 'message': 'Notiz darf nicht leer sein'}), 400

        # Erstelle die Notiz
        note_data = {
            'ticket_id': ObjectId(id),
            'note': note_text,
            'created_by': current_user.username,
            'created_at': datetime.now()
        }
        
        result = mongodb.insert_one('ticket_notes', note_data)
        
        if not result:
            return jsonify({'success': False, 'message': 'Fehler beim Speichern der Notiz'}), 500

        return jsonify({
            'success': True,
            'message': 'Notiz erfolgreich hinzugefügt',
            'note': {
                'id': str(result),
                'text': note_text,
                'created_by': current_user.username,
                'created_at': datetime.now().strftime('%d.%m.%Y %H:%M')
            }
        })

    except Exception as e:
        logging.error(f"Fehler beim Hinzufügen der Notiz: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

def get_unassigned_ticket_count():
    count = mongodb.count_documents('tickets', {
        '$and': [
            {
                '$or': [
                    {'assigned_to': {'$exists': False}},
                    {'assigned_to': None},
                    {'assigned_to': ''}
                ]
            },
            {'status': {'$ne': 'geschlossen'}},
            {'deleted': {'$ne': True}}
        ]
    })
    return count

# Kontextprozessor für alle Templates
@bp.app_context_processor
def inject_unread_tickets_count():
    count = get_unassigned_ticket_count()
    return dict(unread_tickets_count=count)

@bp.route('/auftrag-neu', methods=['GET', 'POST'])
@login_required
def public_create_order():
    """Interne Auftragserstellung für eingeloggte Benutzer."""
    return _handle_auftrag_creation()

@bp.route('/auftrag-extern', methods=['GET', 'POST'])
def external_create_order():
    """Externe Auftragserstellung ohne Login für externe Einbindungen."""
    if request.method == 'POST':
        return _handle_auftrag_creation(external=True)
    categories = get_ticket_categories_from_settings()
    return render_template('tickets/auftrag_external_embed.html', 
                         categories=categories,
                         error=None)

def _handle_auftrag_creation(external=False):
    """Gemeinsame Logik für interne und externe Auftragserstellung."""
    if request.method == 'POST':
        try:
            # Hole die Formulardaten
            title = request.form.get('title')
            description = request.form.get('description')
            category = request.form.get('category')
            priority = request.form.get('priority', 'normal')
            due_date = request.form.get('due_date')
            estimated_time = request.form.get('estimated_time')
            auftraggeber_name = request.form.get('name', '')  # Name des Auftraggebers (Feldname im Template: 'name')
            kontakt = request.form.get('kontakt', '')  # Kontaktdaten
            bereich = request.form.get('bereich', '')  # Bereich
            auftraggeber_typ = request.form.get('auftraggeber_typ', 'extern')  # Intern/Extern
            
            # Validiere die Pflichtfelder
            if not title:
                categories = get_ticket_categories_from_settings()
                if external or not current_user.is_authenticated:
                    return render_template('tickets/auftrag_external_embed.html', 
                                         categories=categories,
                                         error='Titel ist erforderlich.')
                else:
                    flash('Titel ist erforderlich.', 'error')
                    return redirect(url_for('tickets.public_create_order'))
                
            if not description:
                categories = get_ticket_categories_from_settings()
                if external or not current_user.is_authenticated:
                    return render_template('tickets/auftrag_external_embed.html', 
                                         categories=categories,
                                         error='Beschreibung ist erforderlich.')
                else:
                    flash('Beschreibung ist erforderlich.', 'error')
                    return redirect(url_for('tickets.public_create_order'))
                
            # Kategorie ist optional, daher entfernen wir die Validierung
            # if not category:
            #     flash('Kategorie ist erforderlich.', 'error')
            #     return redirect(url_for('tickets.public_create_order'))
            
            # Erstelle das Ticket
            ticket_data = {
                'title': title,
                'description': description,
                'priority': priority,
                'created_by': auftraggeber_name or 'Gast',  # Verwende den eingegebenen Namen oder "Gast" als Fallback
                'category': category,
                'due_date': due_date,
                'estimated_time': estimated_time,
                'status': 'offen',
                'created_at': datetime.now(),
                'updated_at': datetime.now(),
                'ticket_number': get_next_ticket_number(),  # Neue Auftragsnummer
                'is_external': not current_user.is_authenticated  # Markiere externe Aufträge
            }
            
            result = mongodb.insert_one('tickets', ticket_data)
            ticket_id = str(result)
            
            # Erstelle auch die Auftragsdetails mit dem Namen des Auftraggebers
            auftrag_details_data = {
                'ticket_id': ObjectId(ticket_id),
                'bereich': bereich or category,  # Verwende bereich oder category als Fallback
                'auftraggeber_intern': auftraggeber_typ == 'intern',
                'auftraggeber_extern': auftraggeber_typ == 'extern',
                'auftraggeber_name': auftraggeber_name,
                'kontakt': kontakt,
                'auftragsbeschreibung': description,
                'fertigstellungstermin': due_date,
                'gesamtsumme': 0
            }
            
            mongodb.insert_one('auftrag_details', auftrag_details_data)
            
            # Sende Bestätigungs-E-Mail
            if kontakt:  # Nur wenn Kontaktdaten vorhanden sind
                try:
                    from app.utils.email_utils import send_auftrag_confirmation_email
                    # Datum formatieren für E-Mail
                    ticket_data_for_email = ticket_data.copy()
                    if ticket_data_for_email.get('created_at'):
                        ticket_data_for_email['created_at'] = ticket_data_for_email['created_at'].strftime('%d.%m.%Y %H:%M Uhr')
                    
                    send_auftrag_confirmation_email(
                        ticket_data=ticket_data_for_email,
                        auftrag_details=auftrag_details_data,
                        recipient_email=kontakt
                    )
                except Exception as email_error:
                    logger.error(f"Fehler beim Senden der Bestätigungs-E-Mail: {str(email_error)}")
                    # E-Mail-Fehler soll den Auftrag nicht verhindern
            
            # Für alle Benutzer zur Bestätigungsseite
            return render_template('auftrag_public_success.html', 
                                 ticket_number=ticket_data['ticket_number'],
                                 ticket=ticket_data)
            
        except Exception as e:
            logging.error(f"Fehler bei der öffentlichen Auftragserstellung: {str(e)}", exc_info=True)
            categories = get_ticket_categories_from_settings()
            if external or not current_user.is_authenticated:
                return render_template('tickets/auftrag_external_embed.html', 
                                     categories=categories,
                                     error='Ein Fehler ist aufgetreten. Bitte versuchen Sie es später erneut.')
            else:
                flash('Ein Fehler ist aufgetreten. Bitte versuchen Sie es später erneut.', 'error')
                return redirect(url_for('tickets.public_create_order'))
    
    # Hole die Kategorien für das Formular
    categories = get_ticket_categories_from_settings()
    
    if external or not current_user.is_authenticated:
        return render_template('tickets/auftrag_external_embed.html', 
                             categories=categories,
                             error=None)
    return render_template('tickets/create_auftrag.html', 
                         categories=categories,
                         priority_colors={
                             'niedrig': 'secondary',
                             'normal': 'primary',
                             'hoch': 'error',
                             'dringend': 'error'
                         })

@bp.route('/<id>/auftrag-details')
@login_required
def auftrag_details_page(id):
    ticket = mongodb.find_one('tickets', {'_id': ObjectId(id)})
    
    if not ticket:
        return render_template('404.html'), 404
    
    # Prüfe Berechtigungen: Normale User können nur ihre eigenen Tickets sehen
    if current_user.role not in ['admin', 'mitarbeiter'] and ticket.get('created_by') != current_user.username:
        flash('Sie haben keine Berechtigung, dieses Ticket zu sehen.', 'error')
        return redirect(url_for('tickets.index'))
    
    # Füge id-Feld hinzu (für Template-Kompatibilität)
    ticket['id'] = str(ticket['_id'])
    
    # Konvertiere alle Datumsfelder zu datetime-Objekten
    date_fields = ['created_at', 'updated_at', 'resolved_at', 'due_date']
    for field in date_fields:
        if ticket.get(field):
            try:
                ticket[field] = datetime.strptime(ticket[field], '%Y-%m-%d %H:%M:%S')
            except (ValueError, TypeError):
                ticket[field] = None
        
    # Hole die Notizen für das Ticket
    notes = mongodb.find('ticket_notes', {'ticket_id': ObjectId(id)})

    # Hole die Nachrichten für das Ticket
    messages = mongodb.find('ticket_messages', {'ticket_id': ObjectId(id)})

    # Hole die Auftragsdetails
    auftrag_details = mongodb.find_one('auftrag_details', {'ticket_id': ObjectId(id)})
    
    # Hole Materialliste
    material_list = list(mongodb.find('auftrag_material', {'ticket_id': ObjectId(id)}))
    
    # Hole Arbeitsliste
    arbeit_list = list(mongodb.find('auftrag_arbeit', {'ticket_id': ObjectId(id)}))
    
    # Verarbeite die ausgeführten Arbeiten aus den Auftragsdetails
    arbeit_list = []
    if auftrag_details and auftrag_details.get('ausgefuehrte_arbeiten'):
        arbeit_zeilen = auftrag_details['ausgefuehrte_arbeiten'].split('\n')
        for zeile in arbeit_zeilen:
            if zeile.strip():
                teile = zeile.split('|')
                arbeit_list.append({
                    'arbeit': teile[0] if len(teile) > 0 else '',
                    'arbeitsstunden': float(teile[1]) if len(teile) > 1 and teile[1] else 0,
                    'leistungskategorie': teile[2] if len(teile) > 2 else ''
                })
    
    logging.info(f"DEBUG: arbeit_list für Ticket {id}: {arbeit_list}")
    
    # Füge die Auftragsdetails zum Ticket hinzu, damit das Template darauf zugreifen kann
    if auftrag_details:
        ticket['auftrag_details'] = auftrag_details
        # Füge die Material- und Arbeitslisten hinzu
        ticket['auftrag_details']['material_list'] = material_list
        ticket['auftrag_details']['arbeit_list'] = arbeit_list
    else:
        # Falls keine Auftragsdetails vorhanden sind, verwende die Ticket-Daten als Fallback
        ticket['auftrag_details'] = {
            'bereich': ticket.get('category', ''),
            'auftraggeber_name': ticket.get('created_by', ''),
            'auftragsbeschreibung': ticket.get('description', ''),
            'material_list': material_list,
            'arbeit_list': arbeit_list
        }

    return render_template('tickets/auftrag_details_page.html', 
                         ticket=ticket, 
                         notes=notes,
                         messages=messages,
                         auftrag_details=auftrag_details,
                         material_list=material_list,
                         arbeit_list=arbeit_list) 